import pandas as pd
import requests
import os
import folium
import matplotlib.pyplot as plt
import io
import time

from io import BytesIO
from PIL import Image
from folium.plugins import HeatMap
from PIL import Image  
from mecoda_minka import get_obs, get_dfs
from folium.plugins import HeatMap
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime , timedelta
from PIL import Image
import time


# PATHS SETUP

API_PATH = "https://api.minka-sdg.org/v1"


observations = get_obs(id_project=264, grade="research")
df_obs, df_photos = get_dfs(observations)

df_obs.to_csv("data/observations.csv", index=False)
df_photos.to_csv("data/photos.csv", index=False)

# COLOR PALETTES

hex_color_orange = 'f85532'
hex_color_blue = '2b2e4f'

palette_orange = RGBColor.from_string(hex_color_orange)
palette_blue = RGBColor.from_string(hex_color_blue)

########## DATA GENERATION ################

def get_main_metrics(id_project:int):

    """
    Generates a .csv file with the main metrics in the last 30 days
    and the number today

    Parameters:

    - id_project (int) : Number of the project

    Returns :

    - main_metrics.csv
    
    """

    today = datetime.today()
    last_month_date = today - timedelta(days=30)

  
    url_observations_last = f'{API_PATH}/observations?project_id={id_project}&d1={last_month_date}&quality_grade=research'
    observations_results_last = requests.get(url_observations_last).json()['total_results']

    url_observers_last = f'{API_PATH}/observations/observers?project_id={id_project}&d1={last_month_date}&quality_grade=research'
    observers_results_last = requests.get(url_observers_last).json()['total_results']

    url_identifiers_last = f'{API_PATH}/observations/identifiers?project_id={id_project}&d1={last_month_date}&quality_grade=research'
    identifiers_results_last = requests.get(url_identifiers_last).json()['total_results']

    url_species_last = f'{API_PATH}/observations/species_counts?project_id={id_project}&d1={last_month_date}&quality_grade=research'
    species_results_last = requests.get(url_species_last).json()['total_results']


    url_observations_today = f'{API_PATH}/observations?project_id={id_project}&days={today}&quality_grade=research'
    observations_results_today = requests.get(url_observations_today).json()['total_results']

    url_observers_today = f'{API_PATH}/observations/observers?project_id={id_project}&days={today}&quality_grade=research'
    observers_results_today = requests.get(url_observers_today).json()['total_results']

    url_identifiers_today = f'{API_PATH}/observations/identifiers?project_id={id_project}&days={today}&quality_grade=research'
    identifiers_results_today = requests.get(url_identifiers_today).json()['total_results']

    url_species_today = f'{API_PATH}/observations/species_counts?project_id={id_project}&days={today}&quality_grade=research'
    species_results_today = requests.get(url_species_today).json()['total_results']


    df_main_metrics = pd.DataFrame({
        'metric': ['observations', 'observers', 'identifiers', 'species'],
        'number_today': [observations_results_today, observers_results_today, identifiers_results_today, species_results_today],
        'number_in_last_month': [observations_results_last, observers_results_last, identifiers_results_last, species_results_last]
    })

    df_main_metrics.to_csv('data/main_metrics.csv', index=False)
    os.makedirs('data', exist_ok=True)



def get_totals(project_id, year, month, kind="project", session=None):
    if session is None:
        session = requests.Session()

    if kind == "project":
        url_obs = f"{API_PATH}/observations?project_id={project_id}&month={month}&year={year}"
        url_spe = f"{API_PATH}/observations/species_counts?project_id={project_id}&month={month}&year={year}"
        url_part = f"{API_PATH}/observations/observers?project_id={project_id}&month={month}&year={year}"
        url_ident = f"{API_PATH}/observations/identifiers?project_id={project_id}&month={month}&year={year}"
    elif kind == "place":
        url_obs = f"{API_PATH}/observations?place_id={project_id}&month={month}&year={year}"
        url_spe = f"{API_PATH}/observations/species_counts?place_id={project_id}&month={month}&year={year}"
        url_part = f"{API_PATH}/observations/observers?place_id={project_id}&month={month}&year={year}"
        url_ident = f"{API_PATH}/observations/identifiers?place_id={project_id}&month={month}&year={year}"

    total_obs = session.get(url_obs).json().get("total_results", 0)
    total_part = session.get(url_part).json().get("total_results", 0)
    total_ident = session.get(url_ident).json().get("total_results", 0)
    total_spe = session.get(url_spe).json().get("total_results", 0)

    return total_obs, total_part, total_ident, total_spe


def get_month_list(years: list) -> list:
    current_year = datetime.now().year
    current_month = datetime.now().month
    meses = []

    for year in years:
        max_month = 12 if year < current_year else current_month
        for month in range(1, max_month + 1):
            meses.append(f"{year}-{str(month).zfill(2)}")
    return meses

def build_monthly_metrics(project_id):
    meses = get_month_list(range(2022, datetime.now().year + 1))
    total_metrics = []

    session = requests.Session()

    for mes in meses:
        year = int(mes.split("-")[0])
        month = int(mes.split("-")[1])

        print(f"Procesando {mes}...")

        try:
            total_obs, total_part, total_ident, total_spe = get_totals(
                project_id=project_id, year=year, month=month, kind="project", session=session
            )
        except Exception as e:
            print(f"Error procesando {mes}: {e}")
            continue

        total_for_month = {
            "month": mes,
            "observations": total_obs,
            "observers": total_part,
            "identifiers": total_ident,
            "species": total_spe
        }

        total_metrics.append(total_for_month)

    df_monthly = pd.DataFrame(total_metrics)

    os.makedirs("data", exist_ok=True)
    df_monthly.to_csv("data/monthly_metrics.csv", index=False)

    return df_monthly


def get_taxon_count(df_obs:pd.DataFrame, rank_level:str):
    """
    Generates a dataframe with the taxonomy of the project

    Parameters:

    - df_obs (pd.DataFrame): df_obs defined at the begining of the script
    - rank_level (str) : The rank_level selects all the taxonomy in the project

    Returns:

    - df_sorted : Dataframe with three columns ['taxon_rank', 'taxon_name', 'count']
    """

    if rank_level not in df_obs.columns:
        if rank_level == "species":
            df_taxon_counts = df_obs.loc[df_obs['taxon_rank'] == rank_level, "taxon_name"].value_counts().reset_index()
        else:
            raise ValueError(f"'{rank_level}' esta columna no existe en el DataFrame.")

    if rank_level in df_obs.columns:
        df_taxon_counts = df_obs[rank_level].value_counts().reset_index()
        df_taxon_counts.columns = ["taxon_name", "count"]

    if rank_level != "species":
        df2 = df_obs.loc[df_obs['taxon_rank'] == rank_level, "taxon_name"].value_counts().reset_index()

        if len(df2) > 0:
            df_combined = pd.concat([df_taxon_counts, df2])
            df_summed = df_combined.groupby('taxon_name', as_index=False)['count'].sum()
            df_sorted = df_summed.sort_values(by='count', ascending=False).reset_index(drop=True)
            
        else:
            df_sorted = df_taxon_counts
    else:
        df_sorted = df_taxon_counts

    df_sorted['taxon_rank'] = rank_level

    return df_sorted[['taxon_rank', 'taxon_name', 'count']]

def get_taxon_counts():
    
    all_toxon_data = []
    
    for rank_level in ['kingdom', 'phylum', 'class', 'order', 'family', 'genus', 'species']:
        taxon_count = get_taxon_count(df_obs, rank_level)
        all_toxon_data.append(taxon_count)
        print(taxon_count)

    combined_taxon_count = pd.concat(all_toxon_data, ignore_index=True)
    os.makedirs("data", exist_ok=True)
    combined_taxon_count.to_csv("data/taxon_counts.csv", index=False)

def get_new_species(df_obs, last_days=30):
    """
    Obtains the new species registered in the las 30 days

    Parameters: 
    
    - df_obs (pd.DataFrame): df_obs defined at the begining of the script
    - last_days (int): Last 30 days

    Returns:

    - DataFrame with the new species
    """
    df_obs["observed_on"] = pd.to_datetime(df_obs["observed_on"], errors="coerce")

    df_species = df_obs[df_obs["taxon_rank"] == "species"]

    df_species_sorted = df_species.sort_values(by="observed_on")
    df_first = df_species_sorted.drop_duplicates(subset=["taxon_name"], keep="first")

    cutoff_date = datetime.today() - timedelta(days=last_days)
    
    df_new_species = df_first[df_first["observed_on"] > cutoff_date]

    return df_new_species

def get_photos_new_species(df_new_species, df_photos):

    df_photos_new_species = df_photos[df_photos['id'].isin(df_new_species['id'])]
    df_photos_new_species = df_photos_new_species.drop_duplicates(subset=['id'])

    return df_photos_new_species


def download_photos(
    df_photos: pd.DataFrame, directorio: str = "minka_photos"
):
    """
    Function to download the photos resulting from the query.
    """

    if not os.path.exists(directorio):
        os.makedirs(directorio)

    session = requests.Session()
    for i, row in df_photos.iterrows():
        response = session.get(row["photos_medium_url"], stream=True)
        if response.status_code == 200:
            with open(f"{directorio}/{row['path']}", "wb") as out_file:
                out_file.write(response.content)
        del response

    df_photos.loc[:, "abs_path"] = os.path.abspath(f"{directorio}/{df_photos['path']}")


def get_photos_new_species(df_new_species, df_photos):

    df_photos_new_species = df_photos[df_photos['id'].isin(df_new_species['id'])]
    df_photos_new_species = df_photos_new_species.drop_duplicates(subset=['id'])
    download_photos(df_photos_new_species)
    
    df_photos_new_species.to_csv('data/new_species_photos.csv', index=False)

    return df_photos_new_species

def get_new_species_photo_snapshot(df_obs, df_photos, last_days=30):

    '''
    Generates the new_species_photo.csv

    Parameters:

    - df_obs (pd.DataFrame): df_obs defined at the begining of the script
    - df_photos (pd.DataFrame): DataFrame with the photos

    Returns

    - new_species_photo.csv : With all the information of the new_species registered
                              in the last 30 days
    '''

    df_obs["observed_on"] = pd.to_datetime(df_obs["observed_on"], errors="coerce")
    df_species = df_obs[df_obs["taxon_rank"] == "species"]
    df_species_sorted = df_species.sort_values(by="observed_on")
    df_first = df_species_sorted.drop_duplicates(subset=["taxon_name"], keep="first")

    cutoff_date = datetime.today() - timedelta(days=last_days)
    df_new_species = df_first[df_first["observed_on"] > cutoff_date].reset_index(drop=True)


    df_photos_unique = df_photos.drop_duplicates(subset="id")
    photo_map = df_photos_unique.set_index("id")["photos_medium_url"]
    df_new_species["photos_medium_url"] = df_new_species["id"].map(photo_map)

    attribution_map = df_photos_unique.set_index("id")["attribution"]
    df_new_species["attribution"] = df_new_species["id"].map(attribution_map)

    df_new_species['obs_url'] = df_new_species['id'].apply(lambda x: f"https://minka-sdg.org/observations/{x}")

    os.makedirs("data", exist_ok=True)
    
    df_new_species[["taxon_name", "observed_on", "user_login", "photos_medium_url", "attribution", "obs_url"]].to_csv("data/new_species_photo.csv", index=False)

    return df_new_species[["taxon_name", "observed_on", "user_login", "photos_medium_url", "attribution", "obs_url"]].reset_index(drop=True)



############# PLOTS GENERATIONS #############

def plot_monthly_metrics(df_monthly:pd.DataFrame, column: str):
    '''
    Generates the monthly metrics plot for each parameter

    Parameters:
    
    - df_monthly (pd.DataFrame) : DataFrame of the monthly metrics
    - column (str): Each column from the DataFrame can be selected

    Returns

    - Plot for the monthly metrics

    '''
    plt.figure(figsize=(12, 7))
    
    bars = plt.bar(df_monthly['month'], 
                  df_monthly[column], 
                  color='#f85532')
    
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., 
                 height * 1.02,
                 f'{height}',
                 ha='center', 
                 va='bottom')
    
    plt.title(f'Evolució mensual de {column}', pad=20, fontsize=14)
    plt.xlabel('Mes', labelpad=10)
    plt.ylabel(column, labelpad=10)
    
    plt.xticks(rotation=45, ha='right', rotation_mode='anchor') 

    plt.ylim(0, max(df_monthly[column]) * 1.15)

    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['left'].set_visible(False)
    plt.gca().yaxis.set_ticks_position('none')
    
    plt.tight_layout()

    monthly_dir = 'figures/monthly_metrics'
    os.makedirs(monthly_dir, exist_ok=True)
    img_path = os.path.join(monthly_dir, f'monthly_{column}.png')
    plt.savefig(img_path, dpi=300, bbox_inches='tight')
    plt.close()


def plot_top_species(df_taxon_count, taxon_rank: str):

    '''

    Generates the plots for each taxonomic rank

    Parameters:

    - df_taxon_count (pd.DataFrame): Contains all the taxonomic levels from the project
    - taxon_rank (str) : Each taxon rank from the DataFrame can be selected
    
    Returns:

    - Plot of the taxon ranks
    
    '''

    top_df = df_taxon_count[df_taxon_count['taxon_rank'] == taxon_rank]\
             .sort_values('count', ascending=False)\
             .head(10)
    
    plt.figure(figsize=(10, 6))
    bars = plt.barh(top_df['taxon_name'], top_df['count'], color='#f85532')
    
    # Añadir etiquetas con formato
    for bar in bars:
        width = bar.get_width()
        plt.text(width + (0.01 * top_df['count'].max()), 
                bar.get_y() + bar.get_height()/2,
                f'{width:,}',
                va='center')
    
    plt.title(f'Top {taxon_rank} més observats')
    plt.xlabel('Número d\'observacions')
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['left'].set_visible(False)
    plt.gca().yaxis.set_ticks_position('none')  
    plt.gca().invert_yaxis()
    plt.tight_layout()
    
    plot_dir = 'figures/taxon_plots'
    os.makedirs(plot_dir, exist_ok=True)
    img_path = os.path.join(plot_dir, f'top_{taxon_rank}.png')
    plt.savefig(img_path, dpi=300, bbox_inches='tight')
    plt.close()

def get_heatmap(iconic_taxon = None):

    '''
    Contains the Heatmap plots

    Returns:

    - The HeatMap plots 

    '''
    observations_path = 'data/observations.csv'
    df_observations = pd.read_csv(observations_path)
    os.makedirs("figures/heatmap_plots", exist_ok=True)
    df_valid = df_observations.dropna(subset=["latitude", "longitude"])
    
    if iconic_taxon is not None:
        df_valid = df_valid[df_valid["iconic_taxon"] == iconic_taxon]
        if df_valid.empty:
            print(f"No hay datos disponibles para el taxón: {iconic_taxon}")
            #return
    
    heat_data = df_valid[["latitude", "longitude"]].values.tolist()

    mean_lat = df_valid["latitude"].mean()
    mean_lon = df_valid["longitude"].mean()
   

    m = folium.Map(location=[mean_lat, mean_lon], zoom_start=11)
    HeatMap(heat_data).add_to(m)

    os.makedirs("figures", exist_ok=True)

    html_path = f"figures/heatmap_plots/heatmap_{iconic_taxon}.html"
    m.save(html_path)

    img_data = m._to_png(3)
    img = Image.open(io.BytesIO(img_data))
    
    img.save(f'figures/heatmap_plots/heatmap_image_{iconic_taxon}.png')

def plot_new_species(df_new_species):

        image_col = 'photos_medium_url'     
        name_col = 'taxon_name'          
        date_col = 'observed_on'           
        user_col = 'user_login'
        attribution_col = 'attribution'

        plot_dir = 'figures/photos_new_species'
        os.makedirs(plot_dir, exist_ok=True)
    
        counter = 1
   
        for index, row in df_new_species.iterrows():
            
                time.sleep(1)
                
                response = requests.get(row[image_col], timeout=10)
                img = Image.open(BytesIO(response.content))

                fig, ax = plt.subplots(figsize=(5, 5))
                ax.imshow(img)
                ax.axis('off')

                fig.suptitle(f"Nom de l'espècie: {row[name_col]} \n Usuari: {row[user_col]} \n Observat el {row[date_col]}", fontsize = 12, ha='center')
                fig.text(0.5,0.01, f'{row[attribution_col]}', ha='right', fontsize=8, style = 'italic')               
                plt.tight_layout()
                plt.show()

                img_path = os.path.join(plot_dir, f"new_species_{counter:03d}.jpeg")
                fig.savefig(img_path, bbox_inches='tight', dpi=300)
                plt.close()

                attribution_filename = f"new_species_{counter:03d}.txt"
                attribution_path = os.path.join(plot_dir, attribution_filename)
                with open(attribution_path, 'w', encoding='utf-8') as f:
                    f.write(row[attribution_col])

                counter += 1

    ########### HYPERLINK GENERATION  ###############

def add_hyperlink(paragraph, url, text):
    
        """ Creates the hyperlink function for the table in minka_docx """

        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
      
        color = OxmlElement('w:color')
        color.set(qn('w:val'), "0000FF")
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    ############ DOCX GENERATION ################




def get_minka_docx():

    df_main_metrics = pd.read_csv("data/main_metrics.csv")
    df_new_species = pd.read_csv("data/new_species_photo.csv")
    
    last_month = datetime.now().month - 1
    current_year = datetime.now().year

    doc = Document()
      
    doc.add_paragraph()  

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    doc.styles['Normal'].language_id = 1027

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME MENSUAL DEL PROJECTE")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(24)
    run.font.color.rgb = palette_orange

    
    fecha_paragraph = doc.add_paragraph()  

    run_fecha = fecha_paragraph.add_run(F"Informe del {last_month:02d} de {current_year}")  
    run_fecha.font.name = 'Arial Rounded MT Bold' 
    run_fecha.font.size = Pt(18)
    run_fecha.font.color.rgb = palette_blue
    fecha_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    intro = ("En aquest informe es presenten les principals mètriques i estadístiques del projecte. "
        "Aquest document ha estat elaborat a partir de dades de ciència ciutadana "
        "recollides a través de la plataforma web de Minka, una eina col·laborativa per a l’observació "
        "i registre de la biodiversitat.\n"
        "L’objectiu d’aquest informe és oferir una visió general dels resultats obtinguts i facilitar "
        "l’anàlisi de la participació i de les dades observacionals registrades pels usuaris."
    )

    doc.add_paragraph(intro)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Principals mètriques observades l'últim mes")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue


    intro_main_metrics = ("Les dades que es presenten a continuació corresponen a un resum mensual generat de "
    "manera automàtica a partir dels registres disponibles a la plataforma Minka. Aquestes mètriques reflecteixen "
    "l’estat actual del projecte i la seva evolució en els darrers 30 dies pels usuaris."
    )

    doc.add_paragraph(intro_main_metrics)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mètrica'
    hdr_cells[1].text = 'Valor actual'
    hdr_cells[2].text = 'Variació últim mes'

    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
    title_metrics = {
        "observations": "Observacions",
        "observers": "Observadors",
        "identifiers": "Identificadors",
        "species": "Espècies"
    }

    for _, row in df_main_metrics.iterrows():
        row_cells = table.add_row().cells


        metric_ca = title_metrics.get(row['metric'].lower(), row['metric'].capitalize())
        para0 = row_cells[0].paragraphs[0]
        para0.text = metric_ca
        para0.alignment = WD_ALIGN_PARAGRAPH.LEFT

        para1 = row_cells[1].paragraphs[0]
        para1.text = str(row['number_today'])
        para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        para2 = row_cells[2].paragraphs[0]
        para2.text = str(row['number_in_last_month'])
        para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    
    doc.add_page_break()
    
    
    # PLOTS MONTHLY METRICS
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Representació gràfica de les métriques principals de forma mensual")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue
    
    intro_grafics = "A continuació, es presenten els gràfics corresponents a les principals variables analitzades: "
    "Observacions, observadors, identificadors i espècies. Aquests gràfics tenen com a objectiu proporcionar "
    "una visió general  de l'activitat registrada al vostre projecte.\n\n"

    doc.add_paragraph(intro_grafics)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Observacions")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue

    doc.add_picture("figures/monthly_metrics/monthly_observations.png", width=Inches(5.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Observadors")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue

    doc.add_picture("figures/monthly_metrics/monthly_observers.png", width=Inches(5.5))
    
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Identificadors")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue

    doc.add_picture("figures/monthly_metrics/monthly_identifiers.png", width=Inches(5.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Espècies")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue

    doc.add_picture("figures/monthly_metrics/monthly_species.png", width=Inches(5.5))
    
    doc.add_page_break()


    # TAXON PLOTS

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Seguiment de les taxonomies a diferents nivellls")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue
    
    intro_taxon_plots = 'A continuació, es presenten els gràfics corresponents als diferents nivells taxonomics:' \
    'Regne, familia, clase, filo i espècie'

    doc.add_paragraph(intro_taxon_plots)

    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Classificació per regne")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue
   
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/taxon_plots/top_kingdom.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Classificació per clase")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue
   
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/taxon_plots/top_class.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Classificació per filo")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue
   
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/taxon_plots/top_phylum.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Classificació per família")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue

    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/taxon_plots/top_family.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Classificació per espècies")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue
   
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/taxon_plots/top_species.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_page_break()

    # HEATMAP PLOTS
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Mapes de calor")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue

    intro_heatmap = ("Finalment,en aquest apartat es presenta un mapa de calor que mostra la distribució espacial de les observacions registrades. "
        "La densitat d’observacions es representa mitjançant una escala de colors que facilita la detecció de patrons espacials, sent de color vermell les zones amb alta densitat d'observacions" \
        "i les blaves zones amb baixa densitat d'observacions \n\n"
    )

    doc.add_paragraph(intro_heatmap)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Mapa de calor de les observacions totals")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue

    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/heatmap_plots/heatmap_image_None.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    doc.add_page_break()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Mapa de calor de la clase animalia")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue
    
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/heatmap_plots/heatmap_image_animalia.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Mapa de calor de la clase aves")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue
    
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/heatmap_plots/heatmap_image_aves.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    doc.add_page_break()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Mapa de calor de la clase plantae")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue
    
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/heatmap_plots/heatmap_image_plantae.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Mapa de calor de la clase fungi")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(12)
    run.font.color.rgb = palette_blue
     
    paragraph_heatmap = doc.add_paragraph()
    run = paragraph_heatmap.add_run()
    run.add_picture("figures/heatmap_plots/heatmap_image_fungi.png", width=Inches(6))
    paragraph_heatmap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Noves espècies registrades")
    run.font.name = 'Arial Rounded MT Bold'  
    run.font.size = Pt(14)
    run.font.color.rgb = palette_blue

    intro_new_species = (
        "A continuació es presenten totes les espècies novament registrades a la plataforma "
        "durant el període considerat. Feu clic a 'Enllaç' per veure l'observació original."
    )
    doc.add_paragraph(intro_new_species)

  
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Nom de l'espècie"
    hdr_cells[1].text = "Observat al"
    hdr_cells[2].text = "Usuari"
    hdr_cells[3].text = "Enllaç"

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    for _, row in df_new_species.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['taxon_name'])
        row_cells[1].text = str(row['observed_on'])
        row_cells[2].text = str(row['user_login'])
        
        paragraph = row_cells[3].paragraphs[0]
        add_hyperlink(paragraph, row['obs_url'], "Enllaç")



    photo_files = [f for f in os.listdir('figures/minka_photos_new_species') if f.lower().endswith(('.jpeg'))]
    photo_files = photo_files[:5]

    if photo_files:
        doc.add_page_break()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run("Algunes fotografies de les noves espècies registrades")
        run.font.name = 'Arial Rounded MT Bold'
        run.font.size = Pt(14)
        run.font.color.rgb = palette_blue

        for photo_name in photo_files:
            image_path = os.path.join(df_new_species, photo_name)
            
            with Image.open(image_path) as img:
                img = img.convert("RGB")  
                img.thumbnail((800, 800))  
                temp_path = os.path.join(df_new_species, f"resized_{photo_name}")
                img.save(temp_path, format="JPEG", quality=60) 

            paragraph_img = doc.add_paragraph()
            run_img = paragraph_img.add_run()
            run_img.add_picture(temp_path, width=Inches(5.5))
            paragraph_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraf_final = doc.add_paragraph()
    paragraf_final.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_final = paragraf_final.add_run(
    "La resta de fotografies de les noves espècies registrades es poden consultar a la carpeta: figures/photos_new_species"
    )
    run_final.font.name = 'Times New Roman'
    run_final.font.size = Pt(12) 

    doc.save("informe_mensual_minka.docx")

